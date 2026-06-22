"""Optional direction for Word automation."""
from pathlib import Path

Path("reports").mkdir(exist_ok=True)

try:
    import docx  # type: ignore
except ImportError:
    print("如需生成 Word，可安装：python -m pip install python-docx")
else:
    doc = docx.Document()
    doc.add_heading("科研卡片工厂结课报告", level=1)
    doc.add_paragraph("这是由 Python 自动生成的 Word 文档。后面的脚本会把 CSV、Excel、Word 和 PPT 串成一个完整办公自动化包。")
    doc.save("reports/final_report.docx")
    print("已生成 reports/final_report.docx")
